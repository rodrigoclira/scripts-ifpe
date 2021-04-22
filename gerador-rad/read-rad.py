import docx
import csv
import pathlib
from os import path

if __name__ == "__main__":

    # Load the first table from your document. In your example file,
    # there is only one table, so I just grab the first one.
    document = docx.Document('rad.docx')
    table = document.tables[0]
    print(document.tables)
    # Data will be a list of rows represented as dictionaries
    # containing each row's data.
    data = []

    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)

        # Establish the mapping based on the first row
        # headers; these will become the keys of our dictionary
        if i == 0:
            keys = tuple(text)
            continue

        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        data.append(row_data)
    print(data)
