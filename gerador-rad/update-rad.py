import docx
import csv
import pathlib
from os import path

if __name__ == "__main__":

    # Load the first table from your document. In your example file,
    # there is only one table, so I just grab the first one.
    document = docx.Document('rad.docx')
    table_identification = document.tables[0]
    print(document.tables)
    # Data will be a list of rows represented as dictionaries
    # containing each row's data.
    data = []

    keys = None
    text = table_identification.cell(0, 0).text #SIAPE
    text = text.replace("#SIAPE#", "2419898")
    table_identification.cell(0, 0).text  = text
    print(table_identification.cell(0, 1).text) #Coordenação
    print(table_identification.cell(1, 0).text) #Nome
    print(table_identification.cell(1, 1).text) #Grupo
    print(table_identification.cell(1, 2).text) #Regime

    table_ensino = document.tables[1]
    table_ensino.cell(0,1).text = r"""Lorem ipsum dolor sit amet, consectetur adipiscing elit. Nam iaculis at felis ut scelerisque. Praesent ac turpis hendrerit, aliquam nulla in, fermentum leo. Phasellus consequat molestie arcu, pellentesque hendrerit nunc. Nullam vitae libero efficitur, fringilla nibh et, lobortis justo. Curabitur ligula neque, ornare vitae ullamcorper ornare, malesuada in nibh. Maecenas ac imperdiet turpis. Praesent iaculis enim at rhoncus eleifend. Maecenas ut urna facilisis, ornare leo in, venenatis magna. Maecenas vulputate mollis elit, placerat maximus dolor porta in. Morbi sem quam, viverra eget arcu a, mollis vehicula elit.

    Vestibulum placerat enim at lacinia pulvinar. Maecenas tristique nisl dui, at posuere velit iaculis ut. Maecenas at mauris vitae nulla scelerisque blandit. Ut id iaculis quam, sit amet venenatis neque. Nunc eget justo a purus imperdiet hendrerit vel vel lectus. Suspendisse luctus non orci ac consectetur. Interdum et malesuada fames ac ante ipsum primis in faucibus. Ut eu egestas leo, in feugiat neque. Etiam fringilla mi eu interdum convallis.

    Sed at dui eros. Praesent id est nec felis lacinia tristique id nec nisl. Quisque imperdiet quam sed laoreet efficitur. Sed sed ligula varius, cursus orci non, fermentum leo. Proin eu dignissim sapien, non imperdiet ligula. Praesent posuere, erat nec lacinia mollis, est dolor consectetur neque, id accumsan nulla dolor sit amet arcu. Phasellus posuere augue in augue dapibus, ac fermentum tellus suscipit. Maecenas malesuada ligula et ullamcorper venenatis. Vestibulum blandit, magna in finibus vestibulum, mauris arcu aliquet lectus, at fringilla arcu nisl id lectus. Nam dolor erat, placerat ac bibendum ac, faucibus vitae nisi. Sed ante ante, porttitor sit amet commodo eleifend, hendrerit at dolor."""


    table_pesquisa = document.tables[2]
    table_pesquisa.cell(0,1).text = r"""Lorem ipsum dolor sit amet, consectetur adipiscing elit. Nam iaculis at felis ut scelerisque. Praesent ac turpis hendrerit, aliquam nulla in, fermentum leo. Phasellus consequat molestie arcu, pellentesque hendrerit nunc. Nullam vitae libero efficitur, fringilla nibh et, lobortis justo. Curabitur ligula neque, ornare vitae ullamcorper ornare, malesuada in nibh. Maecenas ac imperdiet turpis. Praesent iaculis enim at rhoncus eleifend. Maecenas ut urna facilisis, ornare leo in, venenatis magna. Maecenas vulputate mollis elit, placerat maximus dolor porta in. Morbi sem quam, viverra eget arcu a, mollis vehicula elit.

    Vestibulum placerat enim at lacinia pulvinar. Maecenas tristique nisl dui, at posuere velit iaculis ut. Maecenas at mauris vitae nulla scelerisque blandit. Ut id iaculis quam, sit amet venenatis neque. Nunc eget justo a purus imperdiet hendrerit vel vel lectus. Suspendisse luctus non orci ac consectetur. Interdum et malesuada fames ac ante ipsum primis in faucibus. Ut eu egestas leo, in feugiat neque. Etiam fringilla mi eu interdum convallis.

    Sed at dui eros. Praesent id est nec felis lacinia tristique id nec nisl. Quisque imperdiet quam sed laoreet efficitur. Sed sed ligula varius, cursus orci non, fermentum leo. Proin eu dignissim sapien, non imperdiet ligula. Praesent posuere, erat nec lacinia mollis, est dolor consectetur neque, id accumsan nulla dolor sit amet arcu. Phasellus posuere augue in augue dapibus, ac fermentum tellus suscipit. Maecenas malesuada ligula et ullamcorper venenatis. Vestibulum blandit, magna in finibus vestibulum, mauris arcu aliquet lectus, at fringilla arcu nisl id lectus. Nam dolor erat, placerat ac bibendum ac, faucibus vitae nisi. Sed ante ante, porttitor sit amet commodo eleifend, hendrerit at dolor."""

    table_extensao = document.tables[3]
    table_extensao.cell(0,1).text = r"""Lorem ipsum dolor sit amet, consectetur adipiscing elit. Nam iaculis at felis ut scelerisque. Praesent ac turpis hendrerit, aliquam nulla in, fermentum leo. Phasellus consequat molestie arcu, pellentesque hendrerit nunc. Nullam vitae libero efficitur, fringilla nibh et, lobortis justo. Curabitur ligula neque, ornare vitae ullamcorper ornare, malesuada in nibh. Maecenas ac imperdiet turpis. Praesent iaculis enim at rhoncus eleifend. Maecenas ut urna facilisis, ornare leo in, venenatis magna. Maecenas vulputate mollis elit, placerat maximus dolor porta in. Morbi sem quam, viverra eget arcu a, mollis vehicula elit.

    Vestibulum placerat enim at lacinia pulvinar. Maecenas tristique nisl dui, at posuere velit iaculis ut. Maecenas at mauris vitae nulla scelerisque blandit. Ut id iaculis quam, sit amet venenatis neque. Nunc eget justo a purus imperdiet hendrerit vel vel lectus. Suspendisse luctus non orci ac consectetur. Interdum et malesuada fames ac ante ipsum primis in faucibus. Ut eu egestas leo, in feugiat neque. Etiam fringilla mi eu interdum convallis.

    Sed at dui eros. Praesent id est nec felis lacinia tristique id nec nisl. Quisque imperdiet quam sed laoreet efficitur. Sed sed ligula varius, cursus orci non, fermentum leo. Proin eu dignissim sapien, non imperdiet ligula. Praesent posuere, erat nec lacinia mollis, est dolor consectetur neque, id accumsan nulla dolor sit amet arcu. Phasellus posuere augue in augue dapibus, ac fermentum tellus suscipit. Maecenas malesuada ligula et ullamcorper venenatis. Vestibulum blandit, magna in finibus vestibulum, mauris arcu aliquet lectus, at fringilla arcu nisl id lectus. Nam dolor erat, placerat ac bibendum ac, faucibus vitae nisi. Sed ante ante, porttitor sit amet commodo eleifend, hendrerit at dolor."""

    table_adm = document.tables[4]
    table_adm.cell(0,1).text = r"""Lorem ipsum dolor sit amet, consectetur adipiscing elit. Nam iaculis at felis ut scelerisque. Praesent ac turpis hendrerit, aliquam nulla in, fermentum leo. Phasellus consequat molestie arcu, pellentesque hendrerit nunc. Nullam vitae libero efficitur, fringilla nibh et, lobortis justo. Curabitur ligula neque, ornare vitae ullamcorper ornare, malesuada in nibh. Maecenas ac imperdiet turpis. Praesent iaculis enim at rhoncus eleifend. Maecenas ut urna facilisis, ornare leo in, venenatis magna. Maecenas vulputate mollis elit, placerat maximus dolor porta in. Morbi sem quam, viverra eget arcu a, mollis vehicula elit.

    Vestibulum placerat enim at lacinia pulvinar. Maecenas tristique nisl dui, at posuere velit iaculis ut. Maecenas at mauris vitae nulla scelerisque blandit. Ut id iaculis quam, sit amet venenatis neque. Nunc eget justo a purus imperdiet hendrerit vel vel lectus. Suspendisse luctus non orci ac consectetur. Interdum et malesuada fames ac ante ipsum primis in faucibus. Ut eu egestas leo, in feugiat neque. Etiam fringilla mi eu interdum convallis.

    Sed at dui eros. Praesent id est nec felis lacinia tristique id nec nisl. Quisque imperdiet quam sed laoreet efficitur. Sed sed ligula varius, cursus orci non, fermentum leo. Proin eu dignissim sapien, non imperdiet ligula. Praesent posuere, erat nec lacinia mollis, est dolor consectetur neque, id accumsan nulla dolor sit amet arcu. Phasellus posuere augue in augue dapibus, ac fermentum tellus suscipit. Maecenas malesuada ligula et ullamcorper venenatis. Vestibulum blandit, magna in finibus vestibulum, mauris arcu aliquet lectus, at fringilla arcu nisl id lectus. Nam dolor erat, placerat ac bibendum ac, faucibus vitae nisi. Sed ante ante, porttitor sit amet commodo eleifend, hendrerit at dolor."""

    table_obs = document.tables[5]
    table_obs.cell(0,0).text = r"""Lorem ipsum dolor sit amet, consectetur adipiscing elit. Nam iaculis at felis ut scelerisque. Praesent ac turpis hendrerit, aliquam nulla in, fermentum leo. Phasellus consequat molestie arcu, pellentesque hendrerit nunc. Nullam vitae libero efficitur, fringilla nibh et, lobortis justo. Curabitur ligula neque, ornare vitae ullamcorper ornare, malesuada in nibh. Maecenas ac imperdiet turpis. Praesent iaculis enim at rhoncus eleifend. Maecenas ut urna facilisis, ornare leo in, venenatis magna. Maecenas vulputate mollis elit, placerat maximus dolor porta in. Morbi sem quam, viverra eget arcu a, mollis vehicula elit.

    Vestibulum placerat enim at lacinia pulvinar. Maecenas tristique nisl dui, at posuere velit iaculis ut. Maecenas at mauris vitae nulla scelerisque blandit. Ut id iaculis quam, sit amet venenatis neque. Nunc eget justo a purus imperdiet hendrerit vel vel lectus. Suspendisse luctus non orci ac consectetur. Interdum et malesuada fames ac ante ipsum primis in faucibus. Ut eu egestas leo, in feugiat neque. Etiam fringilla mi eu interdum convallis.

    Sed at dui eros. Praesent id est nec felis lacinia tristique id nec nisl. Quisque imperdiet quam sed laoreet efficitur. Sed sed ligula varius, cursus orci non, fermentum leo. Proin eu dignissim sapien, non imperdiet ligula. Praesent posuere, erat nec lacinia mollis, est dolor consectetur neque, id accumsan nulla dolor sit amet arcu. Phasellus posuere augue in augue dapibus, ac fermentum tellus suscipit. Maecenas malesuada ligula et ullamcorper venenatis. Vestibulum blandit, magna in finibus vestibulum, mauris arcu aliquet lectus, at fringilla arcu nisl id lectus. Nam dolor erat, placerat ac bibendum ac, faucibus vitae nisi. Sed ante ante, porttitor sit amet commodo eleifend, hendrerit at dolor."""

    #print(table_ensino.cell(0,1).text)
    document.save("rad-updated.docx")
